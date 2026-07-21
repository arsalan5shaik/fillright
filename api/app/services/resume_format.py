"""In-place résumé tailoring: edit ONLY the experience bullet text inside the
candidate's own uploaded file (PDF or DOCX), preserving the exact original
formatting of everything else. Falls back (returns None) whenever it can't do
this cleanly, so the caller can render the standard template instead.

DOCX is the high-fidelity path (Word/LibreOffice reflow the paragraph
automatically when the text changes). PDF is best-effort: the original bullet's
text is redacted and the tailored text is written back in the same position,
font family class, size, and colour, shrinking the size just enough to stay
within the original bullet's footprint so nothing below it shifts.
"""

from __future__ import annotations

import io
import re

from app.schemas.resume import ParsedResume
from app.schemas.tailored_resume import TailoredResume

_BULLET_PREFIX = re.compile(r"^[\s•●▪‣⁃∙•\-\*·]+")


def _norm(value: str | None) -> str:
    return re.sub(r"\s+", " ", value or "").strip().lower()


def _strip_bullet(text: str) -> str:
    return _BULLET_PREFIX.sub("", text).strip()


def _bullet_pairs(source: ParsedResume, tailored: TailoredResume) -> list[tuple[str, str]]:
    """(original bullet, tailored bullet) pairs, matched by company+title then
    position - only where the text actually changed."""
    source_by_key = {(_norm(e.company), _norm(e.title)): e for e in source.work_experience}
    pairs: list[tuple[str, str]] = []
    for tentry in tailored.work_experience:
        sentry = source_by_key.get((_norm(tentry.company), _norm(tentry.title)))
        if not sentry:
            continue
        for i, original in enumerate(sentry.bullets):
            if i < len(tentry.bullets) and _norm(original) != _norm(tentry.bullets[i]):
                pairs.append((original, tentry.bullets[i]))
    return pairs


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def _set_paragraph_text(paragraph, text: str) -> None:
    """Replace a paragraph's text while keeping its first run's formatting
    (font, size, bold, colour) and the paragraph's style/indentation."""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _all_paragraphs(container) -> list:
    """Every paragraph in the document, INCLUDING those inside tables (résumés
    very often lay the experience section out in a borderless table, which
    document.paragraphs alone skips - the reason in-place editing looked like a
    no-op)."""
    result = list(container.paragraphs)
    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                result.extend(_all_paragraphs(cell))
    return result


def _tokens(text: str) -> set[str]:
    # Split on non-alphanumerics so "LLM-powered" / "~15" / "hours/week" match
    # "llm powered" / "15" / "hours week" - punctuation differences between the
    # parsed text and the document text otherwise tanked the similarity score.
    return set(re.findall(r"[a-z0-9]+", _strip_bullet(_norm(text))))


def _token_similarity(a: str, b: str) -> float:
    ta, tb = _tokens(a), _tokens(b)
    if not ta or not tb:
        return 0.0
    return len(ta & tb) / len(ta | tb)


def tailor_docx_in_place(original: bytes, source: ParsedResume, tailored: TailoredResume) -> bytes | None:
    try:
        from docx import Document
    except Exception:
        return None

    pairs = _bullet_pairs(source, tailored)
    if not pairs:
        return None

    try:
        document = Document(io.BytesIO(original))
    except Exception:
        return None

    paragraphs = [p for p in _all_paragraphs(document) if p.text and p.text.strip()]

    # Match each original bullet to its paragraph by best token-similarity
    # (robust to the parser having lightly normalized the text). Require EVERY
    # bullet to match confidently - if any doesn't, bail to None so the caller
    # renders the template with fully-tailored bullets rather than leaving a
    # partially- or un-changed file.
    used: set[int] = set()
    plan: list[tuple[int, str]] = []  # (paragraph index, new text)
    for original_bullet, new_bullet in pairs:
        best_idx, best_score = -1, 0.0
        for idx, paragraph in enumerate(paragraphs):
            if idx in used:
                continue
            score = _token_similarity(paragraph.text, original_bullet)
            if score > best_score:
                best_idx, best_score = idx, score
        if best_idx < 0 or best_score < 0.6:
            return None  # couldn't confidently place this bullet - fall back
        used.add(best_idx)
        plan.append((best_idx, new_bullet))

    for idx, new_bullet in plan:
        paragraph = paragraphs[idx]
        prefix_match = _BULLET_PREFIX.match(paragraph.text)
        prefix = prefix_match.group(0) if prefix_match else ""
        _set_paragraph_text(paragraph, prefix + new_bullet)

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
def _int_to_rgb(color: int) -> tuple[float, float, float]:
    return ((color >> 16 & 255) / 255, (color >> 8 & 255) / 255, (color & 255) / 255)


def _map_font(fontname: str) -> str:
    f = (fontname or "").lower()
    bold = "bold" in f or "black" in f or "semibold" in f
    italic = "italic" in f or "oblique" in f
    if any(s in f for s in ("times", "serif", "georgia", "garamond", "minion", "roman")):
        base = "tiro"
        if bold and italic:
            return "tibi"
        if bold:
            return "tibo"
        if italic:
            return "tiit"
        return base
    if any(s in f for s in ("courier", "mono", "consol")):
        return "cobo" if bold else "cour"
    # Helvetica family (covers Arial/Calibri/most sans-serif résumés).
    if bold and italic:
        return "hebi"
    if bold:
        return "hebo"
    if italic:
        return "heit"
    return "helv"


def tailor_pdf_in_place(original: bytes, source: ParsedResume, tailored: TailoredResume) -> bytes | None:
    try:
        import fitz  # PyMuPDF
    except Exception:
        return None

    pairs = _bullet_pairs(source, tailored)
    if not pairs:
        return None

    try:
        doc = fitz.open(stream=original, filetype="pdf")
    except Exception:
        return None

    total_replaced = 0
    try:
        for page in doc:
            lines = _page_lines(fitz, page)
            if not lines:
                continue
            right_edge = max(ln["bbox"].x1 for ln in lines)

            edits = []  # (union_rect, fontname, size, color, new_text)
            used: set[int] = set()
            for original_bullet, new_bullet in pairs:
                target = _strip_bullet(_norm(original_bullet))
                group = _match_line_group(lines, target, used)
                if not group:
                    continue
                first = lines[group[0]]
                union = fitz.Rect(first["bbox"])
                for g in group[1:]:
                    union |= lines[g]["bbox"]
                for g in group:
                    used.add(g)
                edits.append((union, _map_font(first["font"]), first["size"], _int_to_rgb(first["color"]), new_bullet))

            if not edits:
                continue

            for union, _f, _s, _c, _t in edits:
                page.add_redact_annot(union, fill=(1, 1, 1))
            page.apply_redactions()

            for union, fontname, size, color, new_text in edits:
                _write_bullet(fitz, page, union, right_edge, fontname, size, color, new_text)
                total_replaced += 1

        # All-or-nothing: only keep the in-place PDF if EVERY changed bullet was
        # located and rewritten; otherwise fall back to the template (which
        # always has the fully-tailored bullets) rather than ship a file with
        # some bullets still original.
        if total_replaced < len(pairs):
            return None
        return doc.tobytes(deflate=True, garbage=3)
    except Exception:
        return None
    finally:
        doc.close()


def _page_lines(fitz, page) -> list[dict]:
    data = page.get_text("dict")
    lines: list[dict] = []
    for block in data.get("blocks", []):
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            text = "".join(s.get("text", "") for s in spans)
            if not text.strip() or not spans:
                continue
            s0 = spans[0]
            lines.append(
                {
                    "text": text,
                    "bbox": fitz.Rect(line["bbox"]),
                    "font": s0.get("font", ""),
                    "size": s0.get("size", 10.0),
                    "color": s0.get("color", 0),
                }
            )
    lines.sort(key=lambda ln: (round(ln["bbox"].y0), ln["bbox"].x0))
    return lines


def _match_line_group(lines: list[dict], target: str, used: set[int]) -> list[int] | None:
    """Find the consecutive run of lines whose combined text reproduces the
    original bullet `target` (already normalized + bullet-stripped)."""
    for idx, ln in enumerate(lines):
        if idx in used:
            continue
        head = _strip_bullet(_norm(ln["text"]))
        if not head:
            continue
        probe = head[:20]
        if not target.startswith(probe):
            continue
        group = [idx]
        acc = head
        j = idx + 1
        while len(acc) < len(target) - 5 and j < len(lines) and j not in used:
            acc += " " + _norm(lines[j]["text"])
            group.append(j)
            j += 1
        return group
    return None


def _wrapped_line_count(font, text: str, size: float, width: float) -> int:
    space = font.text_length(" ", size)
    lines = 1
    cur = 0.0
    for word in text.split():
        wl = font.text_length(word, size)
        if cur > 0 and cur + space + wl > width:
            lines += 1
            cur = wl
        else:
            cur += (space if cur > 0 else 0) + wl
    return lines


def _write_bullet(fitz, page, union, right_edge, fontname, size, color, new_text) -> None:
    text_out = "•  " + new_text
    width = max(right_edge - union.x0, 40)
    height = union.height
    font = fitz.Font(fontname)
    # Shrink the font just enough that the new text needs no more lines than the
    # original bullet occupied, so nothing below it is pushed down.
    orig_lines = max(1, round(height / (size * 1.16)))
    fit_size = size
    while fit_size > 6.5 and _wrapped_line_count(font, text_out, fit_size, width) > orig_lines:
        fit_size -= 0.5
    # Give the textbox a little vertical slack downward for safe rendering.
    rect = fitz.Rect(union.x0, union.y0 - 1, right_edge + 2, union.y0 + max(height, fit_size * 1.2 * orig_lines) + 2)
    page.insert_textbox(rect, text_out, fontname=fontname, fontsize=fit_size, color=color, align=0)


def tailor_in_place(original: bytes, file_type: str, source: ParsedResume, tailored: TailoredResume) -> bytes | None:
    """Dispatch to the DOCX or PDF editor; returns the edited file bytes in the
    SAME format, or None if in-place editing isn't possible (caller then renders
    the standard template)."""
    if file_type == "docx":
        return tailor_docx_in_place(original, source, tailored)
    if file_type == "pdf":
        return tailor_pdf_in_place(original, source, tailored)
    return None
